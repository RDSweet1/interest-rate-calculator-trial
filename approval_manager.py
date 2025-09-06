"""
Approval Workflow Manager for Interest Rate Calculator GUI Development.

This script manages the approval-based development process, ensuring each
phase is properly reviewed and approved before proceeding to the next.
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from test_runner import TestRunner


class ApprovalManager:
    """Manages the approval workflow for GUI development phases."""
    
    def __init__(self):
        self.approval_file = Path("approval_status.json")
        self.phases = {
            "1": {
                "name": "Basic Application Structure",
                "description": "Tkinter framework, window management, version display, menu structure",
                "dependencies": [],
                "status": "pending"
            },
            "2": {
                "name": "Project Management Interface", 
                "description": "Project CRUD operations, data validation, file management",
                "dependencies": ["1"],
                "status": "pending"
            },
            "3": {
                "name": "Calculation Input Forms",
                "description": "Date pickers, rate inputs, principal amounts, payment management",
                "dependencies": ["2"],
                "status": "pending"
            },
            "4": {
                "name": "Report Generation Interface",
                "description": "Report buttons, progress feedback, file management, notifications",
                "dependencies": ["3"],
                "status": "pending"
            }
        }
        self.load_approval_status()
    
    def load_approval_status(self):
        """Load approval status from file."""
        if self.approval_file.exists():
            try:
                with open(self.approval_file, 'r') as f:
                    data = json.load(f)
                    for phase_id, phase_data in data.items():
                        if phase_id in self.phases:
                            self.phases[phase_id]["status"] = phase_data.get("status", "pending")
                            self.phases[phase_id]["approved_by"] = phase_data.get("approved_by", "")
                            self.phases[phase_id]["approved_date"] = phase_data.get("approved_date", "")
            except Exception as e:
                print(f"Error loading approval status: {e}")
    
    def save_approval_status(self):
        """Save approval status to file."""
        data = {}
        for phase_id, phase_data in self.phases.items():
            data[phase_id] = {
                "status": phase_data["status"],
                "approved_by": phase_data.get("approved_by", ""),
                "approved_date": phase_data.get("approved_date", "")
            }
        
        with open(self.approval_file, 'w') as f:
            json.dump(data, f, indent=2)
    
    def get_current_phase(self) -> str:
        """Get the current phase that can be worked on."""
        for phase_id, phase_data in self.phases.items():
            if phase_data["status"] == "pending":
                # Check if dependencies are met
                deps_met = all(
                    self.phases[dep]["status"] == "approved" 
                    for dep in phase_data["dependencies"]
                )
                if deps_met:
                    return phase_id
        return None
    
    def run_phase_tests(self, phase: str) -> bool:
        """Run tests for a specific phase."""
        runner = TestRunner()
        result = runner.run_phase_tests(phase)
        
        if result["success"]:
            print(f"✅ Phase {phase} tests passed!")
            return True
        else:
            print(f"❌ Phase {phase} tests failed!")
            print(f"Error: {result['stderr']}")
            return False
    
    def generate_approval_package(self, phase: str) -> Path:
        """Generate approval package for a phase."""
        doc = Document()
        
        # Title
        doc.add_heading(f'Phase {phase} - Approval Package', 0)
        
        # Phase Information
        phase_data = self.phases[phase]
        doc.add_heading('Phase Information', level=1)
        doc.add_paragraph(f'Phase: {phase}')
        doc.add_paragraph(f'Name: {phase_data["name"]}')
        doc.add_paragraph(f'Description: {phase_data["description"]}')
        doc.add_paragraph(f'Status: {phase_data["status"]}')
        
        # Dependencies
        doc.add_heading('Dependencies', level=1)
        if phase_data["dependencies"]:
            for dep in phase_data["dependencies"]:
                dep_status = self.phases[dep]["status"]
                doc.add_paragraph(f'Phase {dep}: {dep_status}')
        else:
            doc.add_paragraph('No dependencies')
        
        # Test Results
        doc.add_heading('Test Results', level=1)
        
        # Run tests
        test_passed = self.run_phase_tests(phase)
        
        if test_passed:
            doc.add_paragraph('✅ All tests passed successfully!')
        else:
            doc.add_paragraph('❌ Some tests failed. Review test output.')
        
        # Implementation Details
        doc.add_heading('Implementation Details', level=1)
        
        if phase == "1":
            doc.add_paragraph('• Basic Tkinter application framework')
            doc.add_paragraph('• Window management (bring to front, proper sizing)')
            doc.add_paragraph('• Version information in title bar')
            doc.add_paragraph('• Basic menu structure')
            doc.add_paragraph('• Application icon and styling')
            doc.add_paragraph('• Proper button sizing and visibility')
            
        elif phase == "2":
            doc.add_paragraph('• Project list management interface')
            doc.add_paragraph('• Project creation/editing dialogs')
            doc.add_paragraph('• Project deletion with confirmation')
            doc.add_paragraph('• Project import/export functionality')
            doc.add_paragraph('• Data validation and error handling')
            
        elif phase == "3":
            doc.add_paragraph('• Comprehensive input form for calculation assumptions')
            doc.add_paragraph('• Date picker widgets for billing/as-of dates')
            doc.add_paragraph('• Rate input fields with validation')
            doc.add_paragraph('• Principal amount entry forms')
            doc.add_paragraph('• Payment entry interface with add/edit/delete')
            doc.add_paragraph('• Real-time calculation preview')
            
        elif phase == "4":
            doc.add_paragraph('• Report generation buttons and progress feedback')
            doc.add_paragraph('• Output file management')
            doc.add_paragraph('• Success/error notifications')
            doc.add_paragraph('• Report preview functionality')
            doc.add_paragraph('• File location display and opening')
        
        # Approval Checklist
        doc.add_heading('Approval Checklist', level=1)
        doc.add_paragraph('□ Technical Review: Code quality, architecture, error handling')
        doc.add_paragraph('□ Visual Review: Interface mockups, screenshots, user experience')
        doc.add_paragraph('□ Functional Review: Test results, validation criteria')
        doc.add_paragraph('□ User Acceptance: Manual testing results, workflow validation')
        
        # Approval Section
        doc.add_heading('Approval', level=1)
        doc.add_paragraph('Approved by: _________________________')
        doc.add_paragraph('Date: _________________________')
        doc.add_paragraph('Comments:')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('_' * 50)
        
        # Save approval package
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"phase{phase}_approval_package_{timestamp}.docx"
        report_path = Path("docs/test_reports") / filename
        report_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(report_path))
        
        return report_path
    
    def approve_phase(self, phase: str, approved_by: str = "User") -> bool:
        """Approve a phase."""
        if phase not in self.phases:
            print(f"Invalid phase: {phase}")
            return False
        
        # Check dependencies
        phase_data = self.phases[phase]
        deps_met = all(
            self.phases[dep]["status"] == "approved" 
            for dep in phase_data["dependencies"]
        )
        
        if not deps_met:
            print(f"Phase {phase} dependencies not met!")
            return False
        
        # Approve phase
        self.phases[phase]["status"] = "approved"
        self.phases[phase]["approved_by"] = approved_by
        self.phases[phase]["approved_date"] = datetime.now().isoformat()
        
        self.save_approval_status()
        print(f"✅ Phase {phase} approved by {approved_by}")
        return True
    
    def show_status(self):
        """Show current approval status."""
        print("Interest Rate Calculator - Approval Status")
        print("=" * 50)
        
        for phase_id, phase_data in self.phases.items():
            status_icon = "✅" if phase_data["status"] == "approved" else "⏳"
            print(f"{status_icon} Phase {phase_id}: {phase_data['name']}")
            print(f"   Status: {phase_data['status']}")
            if phase_data["status"] == "approved":
                print(f"   Approved by: {phase_data.get('approved_by', 'Unknown')}")
                print(f"   Approved date: {phase_data.get('approved_date', 'Unknown')}")
            print()
        
        current_phase = self.get_current_phase()
        if current_phase:
            print(f"Next phase to work on: Phase {current_phase}")
        else:
            print("All phases completed!")
    
    def interactive_approval(self):
        """Interactive approval process."""
        while True:
            self.show_status()
            
            print("\nOptions:")
            print("1. Generate approval package for current phase")
            print("2. Approve current phase")
            print("3. Show status")
            print("4. Exit")
            
            choice = input("\nEnter choice (1-4): ").strip()
            
            if choice == "1":
                current_phase = self.get_current_phase()
                if current_phase:
                    print(f"Generating approval package for Phase {current_phase}...")
                    package_path = self.generate_approval_package(current_phase)
                    print(f"Approval package saved to: {package_path}")
                    print("Review the package and approve when ready.")
                else:
                    print("No phase available for approval package generation.")
            
            elif choice == "2":
                current_phase = self.get_current_phase()
                if current_phase:
                    approved_by = input("Enter your name for approval: ").strip()
                    if approved_by:
                        self.approve_phase(current_phase, approved_by)
                    else:
                        print("Approval name required.")
                else:
                    print("No phase available for approval.")
            
            elif choice == "3":
                continue
            
            elif choice == "4":
                break
            
            else:
                print("Invalid choice. Please try again.")


def main():
    """Main function for approval workflow."""
    manager = ApprovalManager()
    
    if len(sys.argv) > 1:
        command = sys.argv[1]
        
        if command == "status":
            manager.show_status()
        elif command == "approve":
            if len(sys.argv) > 2:
                phase = sys.argv[2]
                approved_by = sys.argv[3] if len(sys.argv) > 3 else "User"
                manager.approve_phase(phase, approved_by)
            else:
                print("Usage: python approval_manager.py approve <phase> [approved_by]")
        elif command == "package":
            if len(sys.argv) > 2:
                phase = sys.argv[2]
                package_path = manager.generate_approval_package(phase)
                print(f"Approval package saved to: {package_path}")
            else:
                print("Usage: python approval_manager.py package <phase>")
        else:
            print("Unknown command. Use: status, approve, package, or interactive")
    else:
        # Interactive mode
        manager.interactive_approval()


if __name__ == "__main__":
    main()
